"""Build Weeden Elementary School's landscape weekly learning-plan form.

template weeden-elementary-school-v1

This is intentionally hand-authored from Weeden's approved source document.
The source has a weekly merged Standards/DOK area, section-specific colors,
and a long instructional-sequence row that naturally continues across pages;
those are not representable by the generic one-page codegen schema.
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Twips

_THIS_DIR = str(Path(__file__).resolve().parent)
if _THIS_DIR not in sys.path:
    sys.path.insert(0, _THIS_DIR)
from docx_helpers import (
    fixed_layout,
    set_width,
    shade,
    strip_bloat,
    write_content_lines,
    write_plain,
)

DAY_NAMES = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]
DAY_ABBREVIATIONS = ["M", "T", "W", "R", "F"]
LABEL_W = 2100
DAY_W = 2376
COL_DXA = [LABEL_W] + [DAY_W] * 5

ROW_DEFINITIONS = [
    ("Learning Target/\nEssential Questions", "learning_targets", "E69138"),
    ("Do Now-Bell Ringer", "do_now", "F1C232"),
    ("Vocabulary", "vocabulary", "3D85C6"),
    ("I Do/We Do/You Do", "during", "3D85C6"),
    ("Exit Ticket", "assessment", "FF0000"),
    ("Assessments", "assessment", "A64D79"),
    ("Reteach/Small\nGroups", "reteach_small_groups", "8E7CC3"),
    ("Cross-Curriculum\nConnection", "cross_curricular_connection", "00FF00"),
]


def _write_label(cell, text: str, color: str) -> None:
    write_plain(cell, text, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    shade(cell, color)


def _write_day(cell, day: dict, field: str) -> None:
    if day.get("no_school"):
        write_plain(cell, day.get("title") or "No School", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        return
    value = str(day.get(field, "") or "").strip()
    if field == "during":
        write_content_lines(cell, [(line, False) for line in value.split("\n") if line.strip()] or [("", False)])
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
    else:
        write_plain(cell, value, size=10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def build(data: dict, out_path: str) -> None:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Twips(15840)
    section.page_height = Twips(12240)
    section.top_margin = section.bottom_margin = Twips(540)
    section.left_margin = section.right_margin = Twips(540)

    days = {day["name"]: day for day in data.get("days", [])}
    table = doc.add_table(rows=0, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    fixed_layout(table)

    identity = table.add_row().cells
    for index, width in enumerate(COL_DXA):
        set_width(identity[index], width)
    identity[0].merge(identity[1])
    identity[2].merge(identity[3])
    identity[4].merge(identity[5])
    write_plain(identity[0], "Learning Plans:", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    write_plain(identity[2], str(data.get("course") or ""), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    write_plain(identity[4], f"Week Of:  {data.get('week_of', '')}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)

    headers = table.add_row().cells
    for index, width in enumerate(COL_DXA):
        set_width(headers[index], width)
    shade(headers[0], "FFFFFF")
    for index, day_name in enumerate(DAY_ABBREVIATIONS, start=1):
        write_plain(headers[index], day_name, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
        shade(headers[index], "EAD1DC")

    standards = table.add_row().cells
    for index, width in enumerate(COL_DXA):
        set_width(standards[index], width)
    _write_label(standards[0], "Standard/\nDOK", "E06666")
    standards_cell = standards[1]
    for index in range(2, 6):
        standards_cell.merge(standards[index])
    standard_values = []
    for day_name in DAY_NAMES:
        day = days.get(day_name, {})
        value = str(day.get("standards", "") or "").strip()
        if value and value not in standard_values:
            standard_values.append(value)
    write_content_lines(standards_cell, [(value, False) for value in standard_values] or [("", False)])
    for paragraph in standards_cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(9)

    for label, field, color in ROW_DEFINITIONS:
        row = table.add_row().cells
        for index, width in enumerate(COL_DXA):
            set_width(row[index], width)
        _write_label(row[0], label, color)
        for index, day_name in enumerate(DAY_NAMES, start=1):
            _write_day(row[index], days.get(day_name, {"no_school": True}), field)

    doc.save(out_path)
    strip_bloat(out_path)
