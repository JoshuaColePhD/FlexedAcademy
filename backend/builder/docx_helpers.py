"""Generic python-docx / raw-OOXML helpers shared by every school's builder.

Extracted from build_lesson_plan.py (Florence's own hand-written builder) so a
bug fixed here — several of these functions carry scars from real OOXML
validation failures found the hard way, see each docstring — is fixed for
every school that renders through them, not just the one script it was found
in. Nothing here is template-specific: column widths, shading colors, row
labels, and table shape all live in the caller's own layout spec/script, not
in this module.
"""
from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Twips

CHIP_FILL = "e8eaed"  # Google Docs' own dropdown-chip grey


def shade(cell, hex_fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcpr.append(shd)


def set_width(cell, dxa):
    """python-docx's own width setter already writes a correct
    <w:tcW w:type="dxa"> -- do NOT also append one by hand. Doing so produced
    a duplicate tcW that fails OOXML schema validation (caught 2026-08-03);
    LibreOffice tolerated it but Word can flag such files as needing repair."""
    cell.width = Twips(dxa)


def fixed_layout(table):
    """Force fixed table layout. w:tblLayout must be inserted in its
    schema-mandated position within w:tblPr (before tblCellMar/tblLook), not
    appended at the end -- appending put it after tblLook and failed
    validation (caught 2026-08-03)."""
    tblpr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblpr.insert_element_before(
        layout, "w:tblCellMar", "w:tblLook", "w:tblCaption",
        "w:tblDescription", "w:tblPrChange",
    )


def write_plain(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, size=11):
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold


def write_label(cell, label_text):
    write_plain(cell, label_text, bold=True)


def write_dropdown(cell, selected_text, options, control_id, alias="Dropdown"):
    """Render a dropdown-list content control (w:sdt/w:dropDownList) that
    Google Drive's docx->Doc conversion turns into a real interactive Google
    Docs dropdown chip.

    IMPORTANT -- read before "fixing" this (settled 2026-08-03, see
    build_lesson_plan.py's original docstring for the full investigation):

    This does NOT produce an interactive dropdown chip in Google Docs, and
    nothing can -- Google Docs dropdown chips are creatable ONLY through the
    Google Docs UI by a human, and Google's docx handling never reads this
    markup back (open feature request: issuetracker.google.com/issues/468277653).

    Kept anyway for two real payoffs: it IS a working dropdown in actual
    Microsoft Word, and the chip grey (CHIP_FILL) survives Google's import so
    the cell still visually reads like a chip even though it isn't clickable.

    Structure is a deliberate match to Google's own export -- do not add
    <w:tag> or <w:sdtEndPr/> (Google never emits them), and keep w:sdtPr
    child order as alias, id, dropDownList. control_id must be unique per
    control.
    """
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    sdt = OxmlElement("w:sdt")

    sdt_pr = OxmlElement("w:sdtPr")
    alias_el = OxmlElement("w:alias")
    alias_el.set(qn("w:val"), alias)
    id_el = OxmlElement("w:id")
    id_el.set(qn("w:val"), str(control_id))
    dropdown = OxmlElement("w:dropDownList")
    dropdown.set(qn("w:lastValue"), selected_text)
    for opt in options:
        item = OxmlElement("w:listItem")
        item.set(qn("w:displayText"), opt)
        item.set(qn("w:value"), opt)
        dropdown.append(item)
    sdt_pr.append(alias_el)
    sdt_pr.append(id_el)
    sdt_pr.append(dropdown)

    sdt_content = OxmlElement("w:sdtContent")
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CHIP_FILL)
    shd.set(qn("w:val"), "clear")
    rpr.append(color)
    rpr.append(shd)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = selected_text
    run.append(rpr)
    run.append(t)
    sdt_content.append(run)

    sdt.append(sdt_pr)
    sdt.append(sdt_content)

    para._p.append(sdt)


def write_content_lines(cell, lines_with_bold):
    """lines_with_bold: list of (text, bold). Empty text = blank spacer line."""
    cell.text = ""
    first = True
    for text, bold in lines_with_bold:
        para = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        run = para.add_run(text)
        run.font.size = Pt(11)
        run.font.bold = bold


MINIMAL_STYLES_XML = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
<w:docDefaults>
<w:rPrDefault><w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="22"/></w:rPr></w:rPrDefault>
</w:docDefaults>
<w:style w:type="paragraph" w:default="1" w:styleId="Normal">
<w:name w:val="Normal"/>
</w:style>
<w:style w:type="table" w:styleId="TableGrid">
<w:name w:val="Table Grid"/>
<w:basedOn w:val="TableNormal"/>
<w:tblPr><w:tblBorders>
<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>
<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>
<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>
<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>
<w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>
<w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>
</w:tblBorders></w:tblPr>
</w:style>
<w:style w:type="table" w:default="1" w:styleId="TableNormal">
<w:name w:val="Normal Table"/>
</w:style>
</w:styles>'''

STRIP_PARTS = [
    "word/stylesWithEffects.xml",
    "word/theme/theme1.xml",
    "word/numbering.xml",
    "docProps/thumbnail.jpeg",
    "customXml",
]
STRIP_REL_TYPES = ["stylesWithEffects", "theme", "numbering", "customXml"]
STRIP_CONTENT_TYPE_PARTS = [
    "/word/stylesWithEffects.xml", "/word/theme/theme1.xml",
    "/word/numbering.xml", "/customXml/itemProps1.xml",
]


def strip_bloat(docx_path):
    """Post-process a saved .docx in place: drop python-docx's bundled
    default styles/theme/numbering/thumbnail (unused -- every cell built by
    these helpers uses direct run formatting, not named styles) and replace
    styles.xml with a minimal version covering only what's actually
    referenced (Normal, TableGrid). Shrinks ~40KB to ~8KB with no visual
    change -- verified by rendering both and comparing 2026-08-03."""
    docx_path = Path(docx_path)
    work_dir = docx_path.parent / (docx_path.stem + "_strip_tmp")
    if work_dir.exists():
        shutil.rmtree(work_dir)
    work_dir.mkdir()

    with zipfile.ZipFile(docx_path) as z:
        z.extractall(work_dir)

    for part in STRIP_PARTS:
        full = work_dir / part
        if full.is_dir():
            shutil.rmtree(full)
        elif full.is_file():
            full.unlink()
    theme_dir = work_dir / "word" / "theme"
    if theme_dir.is_dir() and not any(theme_dir.iterdir()):
        theme_dir.rmdir()

    (work_dir / "word" / "styles.xml").write_text(MINIMAL_STYLES_XML, encoding="utf-8")

    ct_path = work_dir / "[Content_Types].xml"
    ct = ct_path.read_text(encoding="utf-8")
    for part in STRIP_CONTENT_TYPE_PARTS:
        ct = re.sub(r'<Override PartName="' + re.escape(part) + r'"[^/]*/>', '', ct)
    ct_path.write_text(ct, encoding="utf-8")

    rels_path = work_dir / "word" / "_rels" / "document.xml.rels"
    rels = rels_path.read_text(encoding="utf-8")
    for rel_type in STRIP_REL_TYPES:
        rels = re.sub(r'<Relationship [^>]*Type="[^"]*' + rel_type + r'"[^>]*/>', '', rels)
    rels_path.write_text(rels, encoding="utf-8")

    # The package-level rels also point at docProps/thumbnail.jpeg, which we
    # just deleted. Leaving that dangling is a CRITICAL validation failure --
    # Word can report the whole file as corrupt (caught 2026-08-03; an earlier
    # version of this function only cleaned word/_rels/document.xml.rels).
    root_rels_path = work_dir / "_rels" / ".rels"
    root_rels = root_rels_path.read_text(encoding="utf-8")
    root_rels = re.sub(
        r'<Relationship [^>]*Target="docProps/thumbnail\.jpeg"[^>]*/>', '', root_rels)
    root_rels_path.write_text(root_rels, encoding="utf-8")

    # python-docx's stock template emits <w:zoom w:val="bestFit"/>; the schema
    # wants a percent attribute. Harmless in practice but it keeps validation
    # clean, so normalise it.
    settings_path = work_dir / "word" / "settings.xml"
    if settings_path.is_file():
        settings_xml = settings_path.read_text(encoding="utf-8")
        if "w:percent" not in settings_xml:
            settings_xml = settings_xml.replace(
                '<w:zoom w:val="bestFit"/>', '<w:zoom w:percent="100"/>')
            settings_path.write_text(settings_xml, encoding="utf-8")

    docx_path.unlink()
    with zipfile.ZipFile(docx_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for f in work_dir.rglob("*"):
            if f.is_file():
                zf.write(f, f.relative_to(work_dir))

    shutil.rmtree(work_dir)
