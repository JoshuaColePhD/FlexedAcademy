"""Turning a school's uploaded lesson-plan template into a structural
analysis an admin can trust — with a lot of checking along the way, on
purpose. This used to be "a human opens the file and hand-writes a builder
script"; the intent here is to hand that human a machine-generated starting
point, not to replace their judgment, so every stage is built to be visibly
wrong when it's wrong rather than silently confident.

Six stages. The first three can end the run with a hard `error` finding
before the next stage ever runs; the last three each independently review
the LLM's output from a different angle, so no single check (and no single
model call) is the only thing standing between a bad read and an admin's
screen:

  1. _validate_persisted_file  — the file on disk is what it claims to be
     (right extension, magic bytes match, non-empty, under the size cap,
     not encrypted) before a single byte of it is parsed as a document.
  2. _extract_*_structure      — deterministic, format-specific extraction
     of headings/tables/fonts/sections. No LLM involved; this is the one
     source of truth everything downstream gets checked against. For a
     scanned-image PDF (near-zero extractable text), this stage falls back
     to OCR (_ocr_pdf_text) when tesseract is installed in this environment
     — self-gating, so an environment without it just gets today's clear
     "can't read a scan" error instead of a crash.
  3. _run_structural_checks    — rule-based sanity checks over the
     extraction itself (empty doc, ragged tables, font explosion, no
     visible structure at all) — the same kind of check calendar_intake.py
     runs on parsed calendar weeks, just for template shape instead of
     dates.
  4. llm.analyze_template_structure + _cross_validate_llm_output — the LLM
     maps the extraction onto named sections, and every section it proposes
     is checked against the extraction it was given: a section whose
     `source_evidence` isn't a literal substring of what the model actually
     saw is dropped, not trusted.
  5. _apply_second_pass_verifier — a SECOND, independently-framed LLM call
     (llm.verify_template_sections) audits the sections that survived stage
     4, specifically checking whether each description is an honest read of
     its evidence rather than just present-in-the-text. Catches a
     plausible-but-wrong mapping stage 4's substring check can't (matching
     text isn't the same as matching meaning). Fails open on its own error —
     an unreachable second opinion keeps stage 4's result, it doesn't cost
     the whole analysis.
  6. _check_coverage — deterministic, no LLM: flags any heading or table the
     extraction found that no surviving section, and no
     `unclear_or_ambiguous` note, ever mentions. Stages 4 and 5 only review
     claims the model DID make; this is the complementary check for what it
     silently didn't.

`analyze_uploaded_template` is the only function other modules should call
for the analysis itself, and never raises — a bad upload becomes a
`status: "failed"` result with findings explaining why, not a 500 in the
middle of onboarding.

`run_and_persist` wraps that with the DB writes every caller needs, INCLUDING
auto-activation: a template that clears every one of the six stages above
with literally zero findings of any severity flips its school straight to
'active', skipping the admin queue entirely (see `_meets_auto_activation_bar`
for exactly how strict that bar is, and why). This exists because the
alternative — every template, however clean, sitting in "Training AI…" until
a human happens to notice the pending queue — was the actual bottleneck in
onboarding a new school, not anything the analysis itself couldn't already
tell you.
"""
from __future__ import annotations

import json
import logging
from dataclasses import dataclass
from datetime import UTC, datetime, timedelta
from pathlib import Path

from . import db, llm, mail
from .config import settings
from .errors import AppError

log = logging.getLogger("flexedacademy.template_intake")

# Deliberately narrower than calendar_intake.SUPPORTED_EXTS: a lesson-plan
# template is always the actual document a school prints/edits, not a link
# or a plain-text export, so there's no txt/md path here.
SUPPORTED_EXTS = {".docx", ".pdf"}

_MAX_PARAGRAPHS = 400   # a blank template with more than this is probably a filled-in packet, not a template
_MAX_PDF_PAGES = 20     # ditto for PDFs
_MIN_TEXT_CHARS = 30    # below this, there's nothing to analyze
_MAX_DISTINCT_FONTS = 8  # more than this suggests copy-pasted, inconsistent formatting
_MAX_PDF_CHARS_SAMPLED_PER_PAGE = 4000  # bound cost on a dense page; font/size survey doesn't need every glyph


@dataclass
class Finding:
    stage: str
    check_name: str
    severity: str  # "info" | "warning" | "error"
    message: str

    def to_dict(self) -> dict:
        return {"stage": self.stage, "check_name": self.check_name, "severity": self.severity, "message": self.message}


class TemplateAnalysisFailed(Exception):
    """A hard error at some stage — carries every finding gathered up to and
    including the one that ended the run, so a 'failed' result still shows
    everything that was checked, not just the last thing."""

    def __init__(self, findings: list[Finding]):
        self.findings = findings
        super().__init__(findings[-1].message if findings else "template analysis failed")


# ---------------------------------------------------------------------------
# Stage 1 — the file is what it claims to be
# ---------------------------------------------------------------------------


def _sniff_ext(head: bytes) -> str:
    """Magic bytes only — never trusts a filename or claimed extension."""
    if head.startswith(b"%PDF-"):
        return ".pdf"
    if head[:4] == b"PK\x03\x04":
        return ".docx"
    return ""


def validate_upload(filename: str, path: Path) -> str:
    """Upload-time boundary check, run on a freshly-spooled temp file before
    it's ever moved to permanent storage: the claimed extension must be
    supported and must match the file's actual magic bytes. Returns the
    validated extension. Kept separate from _validate_persisted_file below —
    that one re-checks a file already trusted enough to live in
    uploads/templates/ (size, non-empty, encryption), which a fresh upload
    hasn't earned yet."""
    ext = Path(filename).suffix.lower()
    if ext not in SUPPORTED_EXTS:
        raise AppError(
            "unsupported_file_type",
            f"Can't accept {ext or 'that file type'} as a lesson-plan template.",
            status=400,
            hint=f"Supported: {', '.join(sorted(SUPPORTED_EXTS))}",
        )
    with open(path, "rb") as fh:
        head = fh.read(8)
    if _sniff_ext(head) != ext:
        raise AppError(
            "file_type_mismatch",
            "That file's contents don't match its extension.",
            status=400,
        )
    return ext


def _validate_persisted_file(path: Path, claimed_ext: str) -> None:
    """Re-validates a file already saved to disk — deliberately independent
    of whatever checks ran during upload, so a re-analysis triggered later
    (e.g. an admin's "re-analyze" button, acting on the stored file path
    alone) is just as defended as the original upload request."""
    if claimed_ext not in SUPPORTED_EXTS:
        raise TemplateAnalysisFailed(
            [Finding("validate", "unsupported_type", "error", f"Unsupported file type {claimed_ext!r}.")]
        )
    if not path.is_file():
        raise TemplateAnalysisFailed(
            [Finding("validate", "file_missing", "error", "The uploaded file is missing from disk.")]
        )
    size = path.stat().st_size
    if size == 0:
        raise TemplateAnalysisFailed([Finding("validate", "empty_file", "error", "The uploaded file is empty.")])
    if size > settings.max_doc_bytes:
        raise TemplateAnalysisFailed(
            [
                Finding(
                    "validate",
                    "file_too_large",
                    "error",
                    f"File is {size // (1024 * 1024)}MB, over the {settings.max_doc_bytes // (1024 * 1024)}MB limit.",
                )
            ]
        )
    with path.open("rb") as fh:
        head = fh.read(8)
    sniffed = _sniff_ext(head)
    if sniffed != claimed_ext:
        raise TemplateAnalysisFailed(
            [
                Finding(
                    "validate",
                    "file_type_mismatch",
                    "error",
                    f"File extension {claimed_ext!r} doesn't match its actual contents.",
                )
            ]
        )


# ---------------------------------------------------------------------------
# Stage 2 — deterministic structural extraction
# ---------------------------------------------------------------------------


def _heading_level(style_name: str) -> int:
    if style_name == "Title":
        return 0
    digits = "".join(ch for ch in style_name if ch.isdigit())
    return int(digits) if digits else 1


def _extract_docx_structure(path: Path) -> dict:
    import docx

    try:
        d = docx.Document(str(path))
    except Exception as e:
        raise TemplateAnalysisFailed(
            [Finding("extract", "docx_unreadable", "error", f"Could not open the .docx file: {e}")]
        ) from e

    headings: list[dict] = []
    paragraph_count = 0
    total_chars = 0
    font_names: set[str] = set()
    font_sizes: set[float] = set()
    bold_runs = italic_runs = 0
    list_paragraphs = 0

    for p in d.paragraphs:
        text = p.text.strip()
        style_name = p.style.name if p.style else "Normal"
        if text:
            paragraph_count += 1
            total_chars += len(text)
        if text and (style_name.startswith("Heading") or style_name == "Title"):
            headings.append({"text": text, "style": style_name, "level": _heading_level(style_name)})
        if "List" in style_name:
            list_paragraphs += 1
        for run in p.runs:
            if run.font.name:
                font_names.add(run.font.name)
            if run.font.size:
                font_sizes.add(round(run.font.size.pt, 1))
            if run.bold:
                bold_runs += 1
            if run.italic:
                italic_runs += 1

    tables: list[dict] = []
    for ti, table in enumerate(d.tables):
        rows = [[c.text.strip() for c in row.cells] for row in table.rows]
        col_counts = {len(r) for r in rows}
        total_chars += sum(len(cell) for row in rows for cell in row)
        tables.append(
            {
                "index": ti,
                "row_count": len(rows),
                "col_counts": sorted(col_counts),
                "header_row": rows[0] if rows else [],
                "sample_rows": rows[1:4],
            }
        )

    sections: list[dict] = []
    for s in d.sections:
        header_text = "".join(pp.text for pp in s.header.paragraphs).strip()
        footer_text = "".join(pp.text for pp in s.footer.paragraphs).strip()
        sections.append(
            {
                "page_width_in": round(s.page_width.inches, 2) if s.page_width else None,
                "page_height_in": round(s.page_height.inches, 2) if s.page_height else None,
                "has_header_text": bool(header_text),
                "has_footer_text": bool(footer_text),
            }
        )

    return {
        "format": "docx",
        "paragraph_count": paragraph_count,
        "total_text_chars": total_chars,
        "headings": headings,
        "tables": tables,
        "page_sections": sections,
        "fonts_used": sorted(font_names),
        "font_sizes_used": sorted(font_sizes),
        "bold_run_count": bold_runs,
        "italic_run_count": italic_runs,
        "list_paragraph_count": list_paragraphs,
        "core_title": (d.core_properties.title or "").strip() or None,
    }


_MAX_OCR_PAGES = 5  # blank templates are 1-3 pages; bound cost on anything larger
_OCR_DPI = 200
_OCR_PAGE_TIMEOUT_S = 20


def _tesseract_available() -> bool:
    import shutil as _shutil

    return _shutil.which("tesseract") is not None


def _ocr_pdf_text(path: Path) -> str:
    """Rasterizes up to _MAX_OCR_PAGES pages (via pdf2image/poppler) and OCRs
    each (via pytesseract/tesseract). Raises on any failure rather than
    returning partial or empty text as if it were a real result — the caller
    decides how a failed OCR attempt gets reported, this just doesn't lie
    about having succeeded."""
    import pdf2image
    import pytesseract

    images = pdf2image.convert_from_path(str(path), dpi=_OCR_DPI, first_page=1, last_page=_MAX_OCR_PAGES)
    return "\n".join(pytesseract.image_to_string(img, timeout=_OCR_PAGE_TIMEOUT_S) for img in images)


def _extract_pdf_structure(path: Path) -> dict:
    import pdfplumber

    try:
        pdf = pdfplumber.open(str(path))
    except Exception as e:
        raise TemplateAnalysisFailed(
            [Finding("extract", "pdf_unreadable", "error", f"Could not open the PDF: {e}")]
        ) from e

    with pdf:
        if getattr(pdf, "is_encrypted", False):
            raise TemplateAnalysisFailed(
                [Finding("extract", "pdf_encrypted", "error", "The PDF is password-protected or encrypted.")]
            )

        pages: list[dict] = []
        tables: list[dict] = []
        font_names: set[str] = set()
        font_sizes: set[float] = set()
        total_chars = 0

        for pi, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            total_chars += len(text)
            for c in page.chars[:_MAX_PDF_CHARS_SAMPLED_PER_PAGE]:
                if c.get("fontname"):
                    font_names.add(c["fontname"])
                if c.get("size"):
                    font_sizes.add(round(c["size"], 1))
            try:
                page_tables = page.extract_tables() or []
            except Exception as e:  # noqa: BLE001 — a broken table detector on one page shouldn't fail the whole doc
                log.warning("pdfplumber table extraction failed on page %d: %s", pi, e)
                page_tables = []
            for t in page_tables:
                rows = [[(cell or "").strip() for cell in row] for row in t]
                col_counts = {len(r) for r in rows}
                tables.append(
                    {
                        "page": pi,
                        "row_count": len(rows),
                        "col_counts": sorted(col_counts),
                        "header_row": rows[0] if rows else [],
                        "sample_rows": rows[1:4],
                    }
                )
            pages.append(
                {
                    "page": pi,
                    "width_in": round(page.width / 72, 2),
                    "height_in": round(page.height / 72, 2),
                    "char_count": len(text),
                }
            )

    # A scanned-image PDF has a real page count and real page dimensions but
    # ~zero extractable text — pdfplumber reads what's literally encoded as
    # text, and a scan has none. Try OCR as a fallback ONLY in that case
    # (never for a normal text PDF — OCR is slower and strictly lossier than
    # reading the real text layer), and only if tesseract is actually
    # installed in this environment.
    ocr_status = "not_needed"
    ocr_text = ""
    if total_chars < _MIN_TEXT_CHARS:
        if not _tesseract_available():
            ocr_status = "not_available"
        else:
            try:
                ocr_text = _ocr_pdf_text(path)
            except Exception as e:  # noqa: BLE001 — a broken OCR attempt reports as a finding, not a crash
                log.warning("OCR fallback failed for %s: %s", path, e)
                ocr_status = "failed"
            else:
                ocr_status = "succeeded" if len(ocr_text.strip()) >= _MIN_TEXT_CHARS else "failed"

    return {
        "format": "pdf",
        "page_count": len(pages),
        "total_text_chars": total_chars if ocr_status != "succeeded" else len(ocr_text),
        "pages": pages,
        "tables": tables,
        "fonts_used": sorted(font_names),
        "font_sizes_used": sorted(font_sizes),
        "ocr_status": ocr_status,
        "ocr_text": ocr_text if ocr_status == "succeeded" else None,
    }


# ---------------------------------------------------------------------------
# Stage 3 — rule-based checks over the extraction, no LLM involved
# ---------------------------------------------------------------------------


def _run_structural_checks(structure: dict) -> list[Finding]:
    findings: list[Finding] = []
    fmt = structure["format"]

    if structure["total_text_chars"] < _MIN_TEXT_CHARS:
        if fmt == "pdf":
            ocr_status = structure.get("ocr_status", "not_needed")
            if ocr_status == "not_available":
                findings.append(
                    Finding(
                        "validate",
                        "empty_or_scanned_document",
                        "error",
                        "Almost no extractable text was found — this looks like a scanned image, and OCR isn't "
                        "available in this environment (the tesseract-ocr binary isn't installed here) to recover it.",
                    )
                )
            elif ocr_status == "failed":
                findings.append(
                    Finding(
                        "validate",
                        "empty_or_scanned_document",
                        "error",
                        "This looks like a scanned image, and an OCR pass over it still couldn't recover enough "
                        "readable text — the scan may be too low-quality or the wrong orientation.",
                    )
                )
            else:
                findings.append(
                    Finding(
                        "validate",
                        "empty_or_scanned_document",
                        "error",
                        "Almost no extractable text was found — this is likely a scanned image rather than a real "
                        "document, which this pipeline cannot read the formatting of.",
                    )
                )
        else:
            findings.append(
                Finding("validate", "empty_document", "error", "The document has almost no readable text.")
            )
        return findings  # nothing further to check meaningfully

    if fmt == "pdf" and structure.get("ocr_status") == "succeeded":
        findings.append(
            Finding(
                "validate",
                "ocr_used",
                "warning",
                "This PDF was a scanned image — text was recovered via OCR, but no font, table, or layout "
                "structure could be detected, only raw text content. Treat the proposed sections as lower "
                "confidence than a native document would produce.",
            )
        )

    if fmt == "docx":
        count = structure["paragraph_count"]
        if count > _MAX_PARAGRAPHS:
            findings.append(
                Finding(
                    "validate",
                    "unusually_large_document",
                    "warning",
                    f"{count} paragraphs is unusually large for a blank template — check it isn't a filled-in "
                    "packet rather than an empty template.",
                )
            )
        has_headings = bool(structure["headings"])
        has_tables = bool(structure["tables"])
        has_emphasis = structure["bold_run_count"] > 0 or structure["italic_run_count"] > 0
        if not (has_headings or has_tables or has_emphasis):
            findings.append(
                Finding(
                    "validate",
                    "no_visible_structure",
                    "error",
                    "No headings, tables, or bold/italic text were found — this template has no identifiable "
                    "structure to map fields onto.",
                )
            )
        prev_level = None
        for h in structure["headings"]:
            if prev_level is not None and h["level"] > prev_level + 1:
                findings.append(
                    Finding(
                        "validate",
                        "heading_hierarchy_gap",
                        "warning",
                        f"Heading level jumps from {prev_level} to {h['level']} at {h['text']!r} — "
                        "the outline may be inconsistent.",
                    )
                )
            prev_level = h["level"]
    else:
        if structure["page_count"] > _MAX_PDF_PAGES:
            findings.append(
                Finding(
                    "validate",
                    "unusually_large_document",
                    "warning",
                    f"{structure['page_count']} pages is unusually large for a blank template.",
                )
            )
        has_tables = bool(structure["tables"])
        has_font_variety = len(structure["font_sizes_used"]) > 1
        # An OCR'd scan has neither by construction — its own "ocr_used"
        # warning above already tells the admin structure detection didn't
        # apply here, so this check would be redundant (and wrong: OCR text
        # existing at all IS the structural signal for that document).
        if structure.get("ocr_status") != "succeeded" and not (has_tables or has_font_variety):
            findings.append(
                Finding(
                    "validate",
                    "no_visible_structure",
                    "error",
                    "No tables and only a single font size were detected — this PDF has no identifiable "
                    "structure to map fields onto.",
                )
            )

    for t in structure["tables"]:
        if len(t["col_counts"]) > 1:
            loc = f"table {t['index']}" if fmt == "docx" else f"table on page {t['page']}"
            findings.append(
                Finding(
                    "validate",
                    "ragged_table",
                    "warning",
                    f"The {loc} has inconsistent row widths ({', '.join(str(c) for c in t['col_counts'])} cells) — "
                    "it may have merged cells the extraction can't represent cleanly.",
                )
            )

    if len(structure["fonts_used"]) > _MAX_DISTINCT_FONTS:
        findings.append(
            Finding(
                "validate",
                "font_explosion",
                "warning",
                f"{len(structure['fonts_used'])} distinct fonts were used — formatting may be inconsistent "
                "throughout the document.",
            )
        )

    return findings


# ---------------------------------------------------------------------------
# Stage 4 — LLM interpretation, then cross-checked against what it was given
# ---------------------------------------------------------------------------


def _summarize_for_llm(structure: dict) -> str:
    """The ONLY input the LLM ever sees for this document — never the raw
    file. Every claim it makes has to trace back to a literal substring of
    this text, which is exactly what _cross_validate_llm_output checks."""
    lines = [f"Format: {structure['format']}"]
    if structure["format"] == "docx":
        lines.append(f"Paragraphs with text: {structure['paragraph_count']}")
        lines.append("Headings (in document order):")
        for h in structure["headings"]:
            lines.append(f"  [{h['style']}] {h['text']}")
        lines.append("Tables:")
        for t in structure["tables"]:
            lines.append(f"  Table {t['index']}: {t['row_count']} rows, header row: {t['header_row']}")
            for row in t["sample_rows"]:
                lines.append(f"    row: {row}")
        lines.append(f"Fonts used: {', '.join(structure['fonts_used']) or '(default only)'}")
        lines.append(f"Bulleted/numbered list paragraphs: {structure['list_paragraph_count']}")
    elif structure.get("ocr_status") == "succeeded":
        # A scanned document has no layout signal at all — the OCR'd text
        # itself is the only thing there is to reason about, so it's what
        # goes in the summary instead of an empty tables/fonts listing.
        lines.append(f"Pages: {structure['page_count']} (scanned image — recovered via OCR, no layout structure)")
        lines.append("OCR'd text:")
        lines.append(structure["ocr_text"] or "")
    else:
        lines.append(f"Pages: {structure['page_count']}")
        lines.append("Tables:")
        for t in structure["tables"]:
            lines.append(f"  Page {t['page']} table: {t['row_count']} rows, header row: {t['header_row']}")
            for row in t["sample_rows"]:
                lines.append(f"    row: {row}")
        lines.append(f"Distinct font sizes on the page(s): {structure['font_sizes_used']}")
    return "\n".join(lines)


def _cross_validate_llm_output(analysis: dict, summary_text: str) -> tuple[dict, list[Finding]]:
    findings: list[Finding] = []
    haystack = summary_text.lower()
    proposed = analysis.get("sections") or []
    verified: list[dict] = []

    for s in proposed:
        evidence = (s.get("source_evidence") or "").strip()
        name = s.get("name") or "(unnamed section)"
        if not evidence:
            findings.append(
                Finding("llm_verify", "missing_evidence", "warning", f"Section {name!r} cited no source evidence.")
            )
            continue
        if evidence.lower() not in haystack:
            findings.append(
                Finding(
                    "llm_verify",
                    "unverified_claim",
                    "warning",
                    f"Section {name!r} cites text not found verbatim in the extracted structure "
                    f"({evidence!r}) — dropped rather than trusted.",
                )
            )
            continue
        verified.append(s)

    dropped = len(proposed) - len(verified)
    if proposed and dropped:
        findings.append(
            Finding(
                "llm_verify",
                "sections_dropped",
                "error" if dropped == len(proposed) else "warning",
                f"{dropped} of {len(proposed)} proposed sections could not be verified and were dropped.",
            )
        )
    if not proposed:
        findings.append(
            Finding("llm_verify", "no_sections_proposed", "error", "The model did not propose any sections at all.")
        )

    confidence = analysis.get("overall_confidence")
    if not isinstance(confidence, (int, float)) or not (0.0 <= confidence <= 1.0):
        findings.append(
            Finding(
                "llm_verify",
                "bad_confidence",
                "warning",
                f"Model reported an out-of-range or missing confidence value: {confidence!r}.",
            )
        )

    if analysis.get("recommended_for_auto_use") and dropped:
        findings.append(
            Finding(
                "llm_verify",
                "confidence_mismatch",
                "warning",
                "Model recommended this template for automatic use despite unverifiable sections — "
                "treat that recommendation with suspicion.",
            )
        )

    unclear = analysis.get("unclear_or_ambiguous") or []
    if unclear:
        findings.append(
            Finding(
                "llm_verify",
                "model_flagged_ambiguity",
                "info",
                f"Model flagged {len(unclear)} unclear/ambiguous item(s) for human review: {unclear}",
            )
        )

    result = dict(analysis)
    result["sections"] = verified
    return result, findings


def _apply_second_pass_verifier(user_id: str, analysis: dict, summary_text: str) -> tuple[dict, list[Finding]]:
    """A second, independent LLM call auditing the sections that already
    survived substring cross-validation — see llm.verify_template_sections'
    own docstring for why this is a skeptical review rather than a re-run of
    the same prompt. Fails open: if the verifier call itself errors (a
    transient LLM/network issue), the sections are kept as-is rather than
    dropping a whole analysis because a SECOND opinion couldn't be reached —
    the first-pass substring check has already run either way."""
    sections = analysis.get("sections") or []
    if not sections:
        return analysis, []

    try:
        verdicts_response = llm.verify_template_sections(user_id, summary_text, sections)
    except Exception as e:  # noqa: BLE001 — see docstring: fail open, not closed
        log.warning("second-pass template verifier unavailable: %s", e)
        return analysis, [
            Finding(
                "llm_verify",
                "second_pass_unavailable",
                "info",
                "The independent verification pass could not be reached; sections were kept based on the "
                "first-pass check alone.",
            )
        ]

    verdicts_by_name = {v.get("name"): v for v in (verdicts_response.get("verdicts") or [])}
    findings: list[Finding] = []
    kept: list[dict] = []
    for s in sections:
        verdict = verdicts_by_name.get(s.get("name"))
        if verdict is None:
            findings.append(
                Finding(
                    "llm_verify",
                    "second_pass_no_verdict",
                    "info",
                    f"Section {s.get('name')!r} got no verdict from the independent verifier — kept, but unconfirmed.",
                )
            )
            kept.append(s)
        elif verdict.get("accurate"):
            kept.append(s)
        else:
            findings.append(
                Finding(
                    "llm_verify",
                    "second_pass_rejected",
                    "warning",
                    f"Independent verifier rejected section {s.get('name')!r}: {verdict.get('reason', '(no reason given)')}",
                )
            )

    result = dict(analysis)
    result["sections"] = kept
    return result, findings


def _check_coverage(structure: dict, analysis: dict) -> list[Finding]:
    """Deterministic complement to _cross_validate_llm_output: that stage
    catches the model claiming something the extraction doesn't support;
    this one catches the opposite mistake — a heading or table the
    extraction actually found that no surviving section, and no
    unclear_or_ambiguous note, ever mentions. That's a silently skipped part
    of the template, not a hallucination, and neither substring
    cross-validation nor the second-pass verifier (which only reviews
    sections the model DID propose) would ever catch it."""
    if structure.get("ocr_status") == "succeeded":
        return []  # OCR text has no headings/tables for coverage to check against

    fmt = structure["format"]
    landmarks: list[tuple[str, str]] = []  # (kind, text)
    if fmt == "docx":
        for h in structure["headings"]:
            landmarks.append(("heading", h["text"]))
    for t in structure["tables"]:
        for cell in t["header_row"]:
            if cell.strip():
                landmarks.append(("table header", cell.strip()))

    covered_text = " ".join(
        [s.get("source_evidence", "") for s in (analysis.get("sections") or [])]
        + list(analysis.get("unclear_or_ambiguous") or [])
    ).lower()

    findings: list[Finding] = []
    seen: set[tuple[str, str]] = set()
    for kind, text in landmarks:
        key = (kind, text.lower())
        if key in seen:
            continue
        seen.add(key)
        if text.lower() not in covered_text:
            findings.append(
                Finding(
                    "coverage",
                    "uncovered_" + kind.replace(" ", "_"),
                    "warning",
                    f"The {kind} {text!r} isn't referenced by any proposed section or listed as unclear — "
                    "it may have been silently missed rather than intentionally excluded.",
                )
            )
    return findings


def _overall_status(findings: list[Finding]) -> str:
    if any(f.severity == "error" for f in findings):
        return "failed"
    if any(f.severity == "warning" for f in findings):
        return "analyzed_with_warnings"
    return "analyzed"


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------


def analyze_uploaded_template(*, user_id: str, dest_path: Path, claimed_ext: str) -> dict:
    """Runs the full pipeline against a template file already saved to disk.

    Never raises: a failure at any stage becomes a `status: "failed"` result
    with findings explaining why, so a bad or ambiguous upload can't 500 the
    endpoint a teacher is sitting in front of mid-onboarding. The file itself
    is always left in place regardless of outcome — analysis is read-only.
    """
    findings: list[Finding] = []
    try:
        _validate_persisted_file(dest_path, claimed_ext)

        structure = _extract_docx_structure(dest_path) if claimed_ext == ".docx" else _extract_pdf_structure(dest_path)
        findings.extend(_run_structural_checks(structure))
        if any(f.severity == "error" for f in findings):
            raise TemplateAnalysisFailed(findings)

        summary_text = _summarize_for_llm(structure)
        try:
            raw_analysis = llm.analyze_template_structure(user_id, summary_text)
        except AppError:
            raise
        except Exception as e:
            raise TemplateAnalysisFailed(
                findings + [Finding("llm", "llm_call_failed", "error", f"Structural interpretation failed: {e}")]
            ) from e

        verified_analysis, verify_findings = _cross_validate_llm_output(raw_analysis, summary_text)
        findings.extend(verify_findings)

        verified_analysis, second_pass_findings = _apply_second_pass_verifier(user_id, verified_analysis, summary_text)
        findings.extend(second_pass_findings)

        findings.extend(_check_coverage(structure, verified_analysis))

        return {
            "status": _overall_status(findings),
            "structure": structure,
            "analysis": verified_analysis,
            "findings": [f.to_dict() for f in findings],
        }
    except TemplateAnalysisFailed as e:
        return {"status": "failed", "structure": None, "analysis": None, "findings": [f.to_dict() for f in e.findings]}
    except Exception as e:
        log.exception("template analysis crashed unexpectedly for %s", dest_path)
        findings.append(Finding("pipeline", "unexpected_error", "error", str(e)))
        return {"status": "failed", "structure": None, "analysis": None, "findings": [f.to_dict() for f in findings]}


# Deliberately conservative — this is the one place in the whole pipeline
# that skips a human ever looking at the template before it goes live for
# real teachers. See _meets_auto_activation_bar's own docstring for why
# every one of these thresholds is stricter than what merely counts as
# "analyzed" elsewhere in this module.
_AUTO_ACTIVATE_MIN_CONFIDENCE = 0.9


def _meets_auto_activation_bar(result: dict) -> bool:
    """True only when EVERY signal agrees the template is unambiguous: not
    just zero errors (that's `status == "analyzed"`), but literally zero
    findings of any severity — including the merely informational ones
    stages 4-6 emit (an unverified second-pass opinion, a model-flagged
    ambiguity) — plus a non-empty section list, and the model's own
    confidence and auto-use recommendation both clearing a high bar.
    Anything short of unanimous leaves the template in the ordinary admin
    queue, exactly as before auto-activation existed."""
    if result["status"] != "analyzed":
        return False
    if result["findings"]:
        return False
    analysis = result["analysis"] or {}
    sections = analysis.get("sections") or []
    if not sections:
        return False
    confidence = analysis.get("overall_confidence")
    if not isinstance(confidence, (int, float)) or confidence < _AUTO_ACTIVATE_MIN_CONFIDENCE:
        return False
    return bool(analysis.get("recommended_for_auto_use"))


def _maybe_auto_activate(school_id: str, template_id: str, result: dict) -> bool:
    """Only ever acts on a school that's still 'pending' — a school that's
    already active (e.g. a re-upload meant to refine an existing format)
    is left alone; that transition stays an admin's call. Returns whether
    it fired, purely so the caller can report it back."""
    school = db.get_school(school_id)
    if not school or school.get("template_status") != "pending":
        return False
    if not _meets_auto_activation_bar(result):
        return False

    db.update_school_template_status(school_id, "active")
    db.mark_template_auto_activated(template_id)
    log.info(
        "auto-activated template %s for school %s (zero findings, confidence>=%.2f)",
        template_id,
        school_id,
        _AUTO_ACTIVATE_MIN_CONFIDENCE,
    )

    latest = db.get_latest_school_template(school_id)
    if latest and latest.get("uploader_email"):
        mail.send_template_active_email(
            to=latest["uploader_email"],
            uploader_name=latest.get("uploader_name"),
            school_name=school["name"],
        )
    return True


def run_and_persist(*, user_id: str, template_id: str, school_id: str, dest_path: Path, claimed_ext: str) -> dict:
    """analyze_uploaded_template plus the DB writes every caller needs around
    it (both the upload endpoint and the admin re-analyze endpoint run this
    exact sequence) — one place so 'mark analyzing, run, save findings, save
    the terminal result, maybe auto-activate' can't drift apart between the
    two call sites."""
    db.set_template_analysis_status(template_id, "analyzing")
    result = analyze_uploaded_template(user_id=user_id, dest_path=dest_path, claimed_ext=claimed_ext)
    db.replace_template_findings(template_id, result["findings"])
    row = db.save_template_analysis(
        template_id,
        status=result["status"],
        structure_json=json.dumps(result["structure"]) if result["structure"] is not None else None,
        analysis_summary=json.dumps(result["analysis"]) if result["analysis"] is not None else None,
        analysis_error=(
            "; ".join(f["message"] for f in result["findings"] if f["severity"] == "error")
            if result["status"] == "failed"
            else None
        ),
    )
    auto_activated = _maybe_auto_activate(school_id, template_id, result)
    if auto_activated:
        row = db.get_school_template(template_id)  # re-fetch: auto_activated flag just changed

    # Kept independent of _maybe_auto_activate above: a slow or failed
    # codegen job (builder_status) must never block or delay the analysis
    # result (template_status) an admin already sees in the pending queue —
    # these are deliberately separate axes, see migration 52's comment.
    # Anything short of a clean or warned analysis with real sections has
    # nothing for a spec to be generated from, so there's no point enqueuing.
    if settings.builder_codegen_enabled and result["status"] in ("analyzed", "analyzed_with_warnings"):
        sections = (result["analysis"] or {}).get("sections") or []
        if sections:
            since = (datetime.now(UTC) - timedelta(days=1)).isoformat(timespec="seconds")
            recent_jobs = db.count_builder_codegen_jobs_created_since(since)
            if recent_jobs < settings.builder_codegen_max_jobs_per_day:
                db.enqueue_builder_codegen_job(school_id, template_id)
            else:
                # Not a failure of THIS upload's analysis — it stays exactly
                # as clean/warned as it already was, in the ordinary admin
                # queue. Just no codegen job for it today; a re-analyze or
                # tomorrow's next upload tries again. Logged, not silently
                # dropped, since "why didn't a job start" should be
                # answerable without reading this function's source.
                log.warning(
                    "builder codegen: daily cap (%d) reached — skipping job for template %s (school %s)",
                    settings.builder_codegen_max_jobs_per_day, template_id, school_id,
                )

    return {"template": row, "findings": result["findings"], "auto_activated": auto_activated}
