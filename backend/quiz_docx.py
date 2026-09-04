"""Render a teacher-facing quiz as a durable Word document.

The quiz JSON remains the canonical structured record.  This renderer is an
export surface, just like qti_build.py: it deliberately produces a simple,
portable DOCX that can be downloaded or converted by Google Drive into an
editable Google Doc.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from . import schema
from .builder.docx_helpers import strip_bloat
from .config import settings
from .docx_build import is_valid_docx, safe_filename

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def quiz_output_path(plan: dict, quiz_id: str) -> Path:
    """Return the stable local path used for a quiz's Word artifact."""
    course = safe_filename(str(plan.get("course") or "Course"), "Course")
    week = safe_filename(str(plan.get("week_of") or "Week"), "Week")
    return Path(settings.plans_dir) / course / "quizzes" / f"{week}__{quiz_id[:8]}.docx"


def _answer_text(question: dict) -> str:
    qtype = question.get("type")
    if qtype == "multiple_choice":
        choices = question.get("choices") or []
        index = question.get("correct_index", -1)
        return f"{chr(65 + index)}. {choices[index]}" if isinstance(index, int) and 0 <= index < len(choices) else ""
    if qtype == "true_false":
        return "True" if question.get("correct_bool") else "False"
    if qtype == "short_answer":
        return "; ".join(str(answer) for answer in (question.get("accepted_answers") or []))
    if qtype == "matching":
        return "; ".join(f"{pair.get('term', '')} — {pair.get('match', '')}" for pair in (question.get("pairs") or []))
    return ""


def _add_question(doc: Document, number: int, question: dict, passages: dict[str, dict] | None = None) -> None:
    passage = (passages or {}).get(question.get("passage_id"))
    if passage and question.get("type") == "multiple_choice":
        passage_title = doc.add_paragraph()
        passage_title.paragraph_format.space_after = Pt(2)
        passage_title.add_run(str(passage.get("title") or "Passage")).bold = True
        passage_text = doc.add_paragraph(str(passage.get("text") or ""))
        passage_text.paragraph_format.left_indent = Inches(0.2)
        passage_text.paragraph_format.right_indent = Inches(0.2)
        passage_text.paragraph_format.space_after = Pt(6)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(f"{number}. {question.get('prompt', '')}")
    run.bold = True

    qtype = question.get("type")
    if qtype == "multiple_choice":
        for index, choice in enumerate(question.get("choices") or []):
            option = doc.add_paragraph(style="List Bullet")
            option.paragraph_format.left_indent = Inches(0.35)
            option.add_run(f"{chr(65 + index)}. {choice}")
    elif qtype == "true_false":
        doc.add_paragraph("True     False")
    elif qtype == "short_answer":
        doc.add_paragraph("________________________________________________________________")
        doc.add_paragraph("________________________________________________________________")
    elif qtype == "matching":
        for pair in question.get("pairs") or []:
            doc.add_paragraph(f"{pair.get('term', '')}  ________")


def build_quiz_docx(quiz: dict, out_path: Path) -> Path:
    """Write a valid DOCX containing the quiz and a teacher answer key."""
    cleaned = schema._clean(quiz)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(str(cleaned.get("title") or "Quiz"))
    title_run.bold = True
    title_run.font.size = Pt(18)

    details = doc.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details.add_run("Name: ________________________________    Date: ________________")

    doc.add_paragraph()
    passages = {
        p.get("id"): p for p in (cleaned.get("passages") or [])
        if isinstance(p, dict) and p.get("id")
    }
    for number, question in enumerate(cleaned.get("questions") or [], 1):
        _add_question(doc, number, question, passages)

    doc.add_section(WD_SECTION.NEW_PAGE)
    key_title = doc.add_paragraph()
    key_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    key_run = key_title.add_run("Teacher Answer Key")
    key_run.bold = True
    key_run.font.size = Pt(16)
    for number, question in enumerate(cleaned.get("questions") or [], 1):
        paragraph = doc.add_paragraph()
        paragraph.add_run(f"{number}. ").bold = True
        paragraph.add_run(_answer_text(question))
        alignment = question.get("alignment") or {}
        if alignment:
            bloom = alignment.get("bloom") or "unassigned"
            dok = alignment.get("dok") or "—"
            cras = (alignment.get("cras") or {}).get("rationale") or ""
            paragraph.add_run(f"  [Bloom: {bloom}; DOK: {dok}; CRAS: {cras}]")

    doc.save(out_path)
    strip_bloat(out_path)
    if not is_valid_docx(out_path):
        raise ValueError("The quiz DOCX renderer produced an invalid Word document.")
    return out_path
