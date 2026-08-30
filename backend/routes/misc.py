"""Settings, chats, health, and file intake."""
from __future__ import annotations

import logging
import re
import subprocess
import tempfile
from collections import Counter
from pathlib import Path

import requests
from fastapi import APIRouter, Depends, File, Request, UploadFile
from fastapi.responses import JSONResponse
from pydantic import BaseModel, Field

from .. import db, docx_build, llm, retrieval, service
from ..config import settings
from ..deps import get_current_user, get_current_user_optional
from ..errors import AppError
from ..ratelimit import limiter

log = logging.getLogger("flexedacademy.misc")
router = APIRouter(prefix="/api", tags=["misc"])

TEXT_EXTS = {".txt", ".md", ".csv"}


# ---------------------------------------------------------------------------
# Health — makes every misconfiguration in the app diagnosable in one curl.
#
# TWO payloads, because it was serving one and it was the detailed one, to
# ANYONE. There is no auth dependency on this route, so a stranger could read
# the model, the on-disk paths of the plans directory and the builder, the
# retrieval thresholds, the corpus size and the length of DATABASE_URL. None of
# that is a credential, but none of it is a stranger's business either, and
# filesystem paths are the first thing anyone probing a host wants.
#
# Liveness has to stay public — Render's health check and any uptime monitor
# call it with no cookie, and a 401 there reads as "the service is down". So
# the public answer is reduced to exactly that: is it up.
# ---------------------------------------------------------------------------


@router.get("/health")
def health(user_id: str | None = Depends(get_current_user_optional)):
    # Unauthenticated: liveness only. Enough for a monitor, useless to a prober.
    if not user_id:
        return {"ok": True}

    out: dict = {
        "ok": True,
        "model": settings.openai_model,
        "api_key_set": settings.has_api_key,
        "database": "PostgreSQL",
        "plans_dir": str(settings.plans_dir),
        "retrieval_floor": settings.retrieval_max_distance,
        "retrieval_top_k": settings.retrieval_top_k,
        # The two settings whose value cannot be inferred from the app's
        # behaviour without either being logged in or trying to break in.
        #
        # require_login=False means every unauthenticated request resolves to
        # 'default_user' — i.e. anyone who loads the site is signed in as the
        # first teacher. That shipped to production once already, and the only
        # way we caught it was noticing /api/chats answered 200 to a stranger.
        # Reporting it turns "is auth on?" from an inference into a fact.
        #
        # Neither is a secret: one is a boolean, the other says whether the
        # session cookie is marked Secure. Knowing them helps a defender far
        # more than an attacker, who can determine both by trying anyway.
        "require_login": settings.require_login,
        # Derived per-request now (see routes/auth.py). Reported as the
        # OVERRIDE flag only — "false" here means "auto", not "insecure".
        "cookie_secure_forced": settings.cookie_secure,
        # Length only, never the value. Every deploy failure so far has been a
        # malformed DATABASE_URL, and each one cost a round trip through the
        # logs to identify. 110 is correct; 160 means the host is in there twice.
        "database_url_len": len(settings.database_url or ""),
    }
    try:
        out["builder_found"] = True
        out["builder_template"] = docx_build.builder_template()
        out["builder_path"] = str(settings.builder_path)
    except AppError as e:
        out["ok"] = False
        out["builder_found"] = False
        out["builder_error"] = e.message
    try:
        out["chunks"] = db._row("SELECT COUNT(*) AS n FROM chunks")["n"]
    except Exception as e:  # noqa: BLE001 — health check reports any DB failure, doesn't crash on one
        out["ok"] = False
        out["pg_error"] = str(e)
    return out


@router.get("/health/ready")
def readiness():
    """Deployment readiness: database reachable and migrations complete.

    Keep this separate from public liveness.  A process can be alive while its
    database pool is unavailable; advertising that process as ready makes a
    deploy route traffic into guaranteed 503s and can hide a failed migration.
    """
    try:
        row = db._row("SELECT COALESCE(MAX(version), 0) AS version FROM schema_version")
        version = int(row["version"]) if row else 0
        expected = len(db.MIGRATIONS)
        if version != expected:
            return JSONResponse(
                {"ok": False, "code": "schema_not_ready", "version": version, "expected": expected},
                status_code=503,
            )
    except Exception:
        log.exception("readiness check failed")
        return JSONResponse({"ok": False, "code": "database_not_ready"}, status_code=503)
    return {"ok": True}


# ---------------------------------------------------------------------------
# Settings — one AP Lang class. Singleton by DB constraint, no profile switcher.
# ---------------------------------------------------------------------------

# Display names for the course codes the ingest scripts write. These ids are what
# the Subject Framework dropdown saves into settings.subject, and what retrieval
# filters on — so this map and scripts/01d_ingest_alcos_case.py FRAMEWORKS have to
# agree. The year is part of the label because Alabama has several live editions
# and a plan citing "Social Studies" is ambiguous without it.
SUBJECT_LABELS = {
    "AP_Lang": "AP Language & Composition",
    "ELA": "English Language Arts (2021)",
    "Math": "Mathematics (2019)",
    "Math_AWF": "Mathematics: Algebra with Finance (2015)",
    "Science": "Science (2023)",
    "Social_Studies": "Social Studies (2024)",
    "Arts": "Arts Education (2024)",
    "DLCS": "Digital Literacy & Computer Science (2025)",
    "Health": "Health Education (2019)",
    "PE": "Physical Education (2019)",
    "World_Languages": "World Languages (2017)",
    "Counseling": "Comprehensive School Counseling (2024-2026)",
    "Special_Education": "Special Education / Collaborative",
}

# AP Lang first: it is the course this app exists for and the default settings row.
_FRAMEWORK_PRIORITY = ("AP_Lang", "ELA")


def known_course_ids() -> set[str]:
    """Every course code a class's `subject` (post service.subject_code)
    can resolve to without breaking retrieval — the same set /api/frameworks
    itself is built from, plus the synthetic "Special_Education" aggregate
    that route always offers regardless of chunk counts (see get_frameworks
    below).

    Shared with routes/classes.py's create/update validation: without this
    check, ClassBody.subject/ClassPatch.subject accept any string, and the
    only thing that ever kept garbage out was FrameworkPicker.jsx offering
    nothing but real course codes — a frontend convention, not a backend
    guarantee. A class created with an unresolvable subject doesn't fail
    loud until the teacher tries to generate a plan for it.
    """
    ids = {c.get("course") for c in retrieval.load_chunks() if c.get("course")}
    ids.add("Special_Education")
    return ids


@router.get("/frameworks")
def get_frameworks():
    """What the Subject Framework and Grade Level dropdowns are built from.

    Derived from the chunks, so a framework is offered only while it is actually
    ingested — the UI can't present a subject retrieval would fail to ground.
    """
    grades: dict[str, set[int]] = {}
    counts: Counter = Counter()
    verbatim: Counter = Counter()

    for c in retrieval.load_chunks():
        course = c.get("course")
        if not course:
            continue
        counts[course] += 1
        if c.get("verbatim_ok"):
            verbatim[course] += 1
        if course.startswith(("AP", "Pre-AP")):
            grades.setdefault(course, set()).update([9, 10, 11, 12])
        else:
            grade = c.get("grade")
            # `if course and grade` was the old test, which dropped every
            # Kindergarten chunk: grade 0 is falsy. It also admitted 99, a
            # "grade unknown" sentinel from an earlier ingest that no teacher can
            # select and that matches nothing when used as a filter.
            if isinstance(grade, int) and 0 <= grade <= 12:
                grades.setdefault(course, set()).add(grade)

    result = [
        {
            "id": course,
            "label": SUBJECT_LABELS.get(course, course.replace("_", " ")),
            "grades": sorted(grades.get(course, set())),
            "chunks": counts[course],
            "verbatim_ok": verbatim[course],
        }
        for course in counts
    ]
    
    if "Special_Education" not in counts:
        core_subjects = ["AP_Lang", "ELA", "Math", "Math_AWF", "Science", "Social_Studies"]
        sp_chunks = sum(counts.get(s, 0) for s in core_subjects)
        sp_verb = sum(verbatim.get(s, 0) for s in core_subjects)
        result.append({
            "id": "Special_Education",
            "label": SUBJECT_LABELS["Special_Education"],
            "grades": [9, 10, 11, 12],
            "chunks": sp_chunks,
            "verbatim_ok": sp_verb,
        })

    result.sort(key=lambda f: (
        _FRAMEWORK_PRIORITY.index(f["id"]) if f["id"] in _FRAMEWORK_PRIORITY
        else len(_FRAMEWORK_PRIORITY),
        f["label"],
    ))
    return result

class SettingsBody(BaseModel):
    teacher: str = Field(min_length=1, max_length=120)
    course: str = Field(min_length=1, max_length=160)
    period: str = Field(min_length=1, max_length=80)
    subject: str = Field(default="AP Language & Composition", max_length=120)
    grade: str = Field(default="11", max_length=20)


@router.get("/settings")
def get_settings_route(subject: str | None = None, user_id: str = Depends(get_current_user)):
    """The settings row, with `subject` resolved to a live course code.

    Normalised on the way out rather than migrated in place, so a row written
    before the frameworks were renamed still drives the UI correctly: the Grade
    Level dropdown is built by looking `subject` up in /api/frameworks, and an
    id that no longer exists there left it stuck on a single hardcoded option.
    The next save persists the resolved value.
    """
    row = dict(db.get_settings_row(user_id, subject))
    row["subject"] = service.subject_code(row.get("subject", ""))
    return row


@router.put("/settings")
def put_settings(body: SettingsBody, user_id: str = Depends(get_current_user)):
    return db.update_settings(
        user_id,
        body.teacher.strip(),
        body.course.strip(),
        body.period.strip(),
        body.subject.strip(),
        body.grade.strip()
    )


# ---------------------------------------------------------------------------
# Chats
# ---------------------------------------------------------------------------


class ChatBody(BaseModel):
    title: str = Field(min_length=1, max_length=200)
    # Optional so an older client still creates a working (unscoped) chat.
    class_id: str | None = Field(default=None, max_length=64)
    # The week this conversation is about, pinned at creation — see db.py's
    # migration 24. Also optional: a chat with no week still works, it just
    # doesn't get the "currently working on …" block in its system prompt.
    # Ignored by the rename route below, which shares this model.
    week_number: int | None = Field(default=None, ge=1, le=52)
    mode: str | None = Field(default=None, max_length=20)


class TitleRequest(BaseModel):
    message: str = Field(min_length=1, max_length=4000)


@router.post("/chats/title")
def suggest_chat_title(body: TitleRequest, user_id: str = Depends(get_current_user)):
    """A short descriptive title generated from a chat's first message.

    Called after the chat is already created with a truncated placeholder
    title, so a slow or failed title suggestion never blocks sending the
    first message. Requires login like every other route here, even though
    it doesn't touch the requester's data — an anonymous caller shouldn't be
    able to spend the shared OpenAI budget.
    """
    return {"title": llm.generate_chat_title(user_id, body.message)}


@router.get("/chats")
def list_chats(class_id: str | None = None, user_id: str = Depends(get_current_user)):
    """Scoped to one prep when class_id is given. Omitting it returns
    everything, which is what /api/chats meant before and what any caller
    without a class in hand still wants."""
    return db.list_chats(user_id, class_id=class_id)


@router.post("/chats")
def create_chat(body: ChatBody, user_id: str = Depends(get_current_user)):
    return db.create_chat(
        user_id, body.title, class_id=body.class_id, week_number=body.week_number, mode=body.mode
    )


@router.get("/chats/{chat_id}")
def get_chat(chat_id: str, user_id: str = Depends(get_current_user)):
    chat = db.get_chat(user_id, chat_id, with_messages=True)
    if not chat:
        raise AppError("chat_not_found", "No such chat.", status=404)
    return chat


@router.patch("/chats/{chat_id}")
def rename_chat(chat_id: str, body: ChatBody, user_id: str = Depends(get_current_user)):
    chat = db.rename_chat(user_id, chat_id, body.title)
    if not chat:
        raise AppError("chat_not_found", "No such chat.", status=404)
    return chat

class PinBody(BaseModel):
    is_pinned: bool

@router.patch("/chats/{chat_id}/pin")
def pin_chat(chat_id: str, body: PinBody, user_id: str = Depends(get_current_user)):
    chat = db.toggle_chat_pin(user_id, chat_id, body.is_pinned)
    if not chat:
        raise AppError("chat_not_found", "No such chat.", status=404)
    return chat


class ChatWeekBody(BaseModel):
    week_number: int = Field(ge=1, le=52)


class ChatModeBody(BaseModel):
    mode: str = Field(pattern=r"^(brainstorm|build|research|interview|standards|sub_plan)$")


@router.patch("/chats/{chat_id}/mode")
def set_chat_mode(chat_id: str, body: ChatModeBody, user_id: str = Depends(get_current_user)):
    chat = db.set_chat_mode(user_id, chat_id, body.mode)
    if not chat:
        raise AppError("chat_not_found", "No such chat.", status=404)
    return chat


@router.patch("/chats/{chat_id}/week")
def set_chat_week(chat_id: str, body: ChatWeekBody, user_id: str = Depends(get_current_user)):
    """Re-point an existing conversation at a different week — the composer's
    own week dropdown, once a chat exists to re-point.

    Its own route rather than a field on the rename PATCH above: that one
    requires a title, which a week change has no business inventing."""
    chat = db.set_chat_week(user_id, chat_id, body.week_number)
    if not chat:
        raise AppError("chat_not_found", "No such chat.", status=404)
    return chat


@router.delete("/chats/{chat_id}", status_code=204)
def delete_chat(chat_id: str, user_id: str = Depends(get_current_user)):
    if not db.delete_chat(user_id, chat_id):
        raise AppError("chat_not_found", "No such chat.", status=404)


# ---------------------------------------------------------------------------
# File intake
# ---------------------------------------------------------------------------


def _spool(upload: UploadFile, suffix: str, max_bytes: int) -> Path:
    """Stream to a temp file, aborting past the cap.

    Enforced while reading rather than from content-length (spoofable) — and the
    old code did `await audio.read()`, buffering an unbounded upload in RAM. The
    filename is generated, never taken from upload.filename, which was
    previously interpolated straight into a path.
    """
    total = 0
    try:
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as fd:
            name = fd.name
            while chunk := upload.file.read(1 << 20):
                total += len(chunk)
                if total > max_bytes:
                    raise AppError(
                        "file_too_large",
                        f"That file is larger than the {max_bytes // (1024 * 1024)}MB limit.",
                        status=413,
                    )
                fd.write(chunk)
    except Exception:
        Path(name).unlink(missing_ok=True)
        raise
    return Path(name)


def _download_google_doc_as_docx(url: str, max_bytes: int) -> Path | None:
    """Download a public Google Doc directly as a .docx file. Returns the Path
    to a temporary file containing the bytes, or raises AppError if it's
    private or fails."""
    # Pattern to extract the document ID from a Google Docs URL
    match = re.search(r"/document/d/([a-zA-Z0-9-_]+)", url)
    if not match:
        return None

    doc_id = match.group(1)
    export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=docx"

    try:
        # A timeout prevents hanging on Google's servers; stream=True prevents
        # buffering a giant file in memory all at once.
        resp = requests.get(export_url, timeout=10, stream=True)
    except requests.RequestException as e:
        log.warning("failed to fetch google doc %s: %s", doc_id, e)
        raise AppError("bad_request", "Failed to connect to Google Docs. Please check the link and try again.", status=400)

    # A private document redirects to a Google login page, which returns 200 OK
    # but with a text/html content type instead of the requested format.
    content_type = resp.headers.get("content-type", "")
    if resp.status_code != 200 or "text/html" in content_type:
        raise AppError(
            "bad_request",
            "This Google Doc is private. Please click 'Share' in Google Docs, set General Access to 'Anyone with the link', and try again.",
            status=400,
        )

    total = 0
    try:
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as fd:
            name = fd.name
            for chunk in resp.iter_content(chunk_size=1 << 20):
                if chunk:
                    total += len(chunk)
                    if total > max_bytes:
                        raise AppError(
                            "file_too_large",
                            f"That Google Doc is larger than the {max_bytes // (1024 * 1024)}MB limit.",
                            status=413,
                        )
                    fd.write(chunk)
    except Exception:
        Path(name).unlink(missing_ok=True)
        raise
    return Path(name)


# 60/minute, not 30: voice mode sends one of these per spoken utterance, and a
# real back-and-forth ("week seven" — "no, the Poe one" — "make Thursday a
# seminar") is a dozen short turns a minute before anyone is being unreasonable.
# The actual spend ceiling is entitlement.py's token cap, which llm.transcribe
# already charges against; this limit is abuse protection, and at 30 it was
# tight enough to fire on ordinary use.
@router.post("/transcribe")
@limiter.limit("60/minute")
def transcribe(request: Request, audio: UploadFile = File(...), user_id: str = Depends(get_current_user)):
    """Was unauthenticated — every other OpenAI-backed route here requires
    login (see /tts's own comment), but this one didn't, which made it a free
    relay for anyone who found it: no account to attribute the cost to, and
    nothing for entitlement.py's cap to apply against.

    Deliberately a plain `def`, NOT `async def` — and this is the whole reason
    voice mode used to stall mid-conversation. Both calls in the body block:
    _spool() does synchronous file I/O, and llm.transcribe() makes a blocking
    OpenAI HTTP request that takes most of a second. Inside an `async def`,
    FastAPI runs the handler ON the event loop, so for the entire Whisper round
    trip the whole server was frozen — the SSE chat stream couldn't push a
    chunk, queued /tts fetches couldn't be served, nothing else could run. On a
    single-worker uvicorn (see render.yaml's free plan) that is every request,
    not just this one. A plain `def` is dispatched to the threadpool instead,
    which is what every other route in this file already does; this one was the
    outlier, and it happened to be the hottest route in the feature.
    """
    suffix = Path(audio.filename or "recording.webm").suffix or ".webm"
    path = _spool(audio, suffix, settings.max_audio_bytes)
    try:
        return {"text": llm.transcribe(user_id, str(path))}
    finally:
        path.unlink(missing_ok=True)


def read_text_from_path(path: Path, ext: str) -> str:
    """Shared by /extract_text and curriculum-map ingestion, so PDF/docx handling
    lives in exactly one place."""
    if ext == ".pdf":
        try:
            out = subprocess.run(
                ["pdftotext", "-layout", str(path), "-"],
                capture_output=True,
                text=True,
                timeout=30,
                check=False,
            )
        except FileNotFoundError as e:
            raise AppError(
                "pdftotext_missing",
                "The pdftotext tool isn't installed.",
                hint="Install poppler: brew install poppler",
            ) from e
        except subprocess.TimeoutExpired as e:
            raise AppError("pdf_timeout", "That PDF took too long to read.", status=422) from e
        if out.returncode != 0:
            raise AppError(
                "pdf_unreadable",
                "Could not extract text from that PDF.",
                status=422,
                hint=(out.stderr or "").strip()[:200] or None,
            )
        return out.stdout
    if ext == ".docx":
        import docx as _docx

        d = _docx.Document(str(path))
        parts = [p.text for p in d.paragraphs if p.text.strip()]
        for table in d.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    return path.read_text(encoding="utf-8", errors="ignore")


@router.post("/extract_text")
def extract_text(file: UploadFile = File(...)):
    # Plain `def` for the same reason /transcribe is — see its docstring. This
    # one was arguably worse while it lasted: read_text_from_path shells out to
    # pdftotext with a 30-second timeout, so a slow PDF could hold the event
    # loop (and therefore every other request on the instance) for half a
    # minute. It just fires far less often than /transcribe does.
    original = Path(file.filename or "upload")
    ext = original.suffix.lower()
    if ext != ".pdf" and ext not in TEXT_EXTS:
        raise AppError(
            "unsupported_file_type",
            f"Can't read {ext or 'that file type'}.",
            status=400,
            hint="Supported: .pdf, .txt, .md, .csv",
        )

    path = _spool(file, ext, settings.max_doc_bytes)
    try:
        text = read_text_from_path(path, ext)
        # Keep the ORIGINAL name for display; it never touched the filesystem.
        return {"filename": original.name, "text": text, "chars": len(text)}
    finally:
        path.unlink(missing_ok=True)


def purge_legacy_temp() -> None:
    """temp/ is no longer a store. Clear the intermediates it used to accumulate.

    Generated .docx files in there are left alone — they may be the only copy of
    a real week, and they predate the plans/ directory and the database.
    """
    temp_dir = Path(settings.plans_dir).parent / "temp"
    if not temp_dir.is_dir():
        return
    removed = 0
    for pattern in ("temp_*.json", "upload_*"):
        for p in temp_dir.glob(pattern):
            try:
                p.unlink()
                removed += 1
            except OSError as e:
                log.warning("could not remove %s: %s", p, e)
    leftovers = sorted(temp_dir.glob("LessonPlan_*.docx"))
    if removed:
        log.info("purged %d legacy temp intermediates", removed)
    if leftovers:
        log.info(
            "left %d pre-existing .docx in temp/ untouched (not in the plans library): %s",
            len(leftovers),
            ", ".join(p.name for p in leftovers),
        )
