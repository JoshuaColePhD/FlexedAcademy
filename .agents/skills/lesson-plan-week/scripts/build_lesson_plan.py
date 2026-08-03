#!/usr/bin/env python3
"""Build a Florence City Schools weekly lesson plan .docx (template florence-docx-v1).

Usage:
    python3 build_lesson_plan.py week.json "Week_11_Oct_19-23.docx"

Input JSON schema (see example-week.json):
{
  "teacher": "Josh Cole",
  "grade_subject": "Grade 11 / AP Language & Composition",
  "week_of": "Week 11 — Oct 19-23, 2026",
  "days": [                       # exactly 5: Monday..Friday
    {
      "name": "Monday",
      "no_school": false,
      "curriculum": "...",                # Row 3
      "standards": "2.A, 2.B",            # Row 4 (prefixed 'Standard:')
      "learning_targets": "I can ...",    # Row 5 (prefixed 'Learning Targets:')
      "do_now": "...",                    # Row 6 (prefixed 'Do Now:')
      "learning_looks_like": "Explicit: ...\nStudent: ...\nIndependent: ...",  # Row 7
      "assessment": "Formative — ..."     # Row 8
    }, ... ]
}

The script reproduces the master template exactly: landscape letter, 0.5" margins,
1980/2484 DXA column widths, header blue (A4C2F4) rows 1-2, label blue (CFE2F3)
rows 3-8 label column, Arial Bold 11 headers, Calibri content. Requires python-docx
(`pip install python-docx`).
"""

import json
import sys

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, Twips

BLUE_HDR = "A4C2F4"   # header rows 1-2
BLUE_LBL = "CFE2F3"   # label column rows 3-8
WHITE = "FFFFFF"
BLACK = "000000"

LABEL_W = 1980        # DXA
DAY_W = 2484          # DXA
COL_DXA = [LABEL_W] + [DAY_W] * 5

DAY_NAMES = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]

# (label, json_key, guidance) for content rows 3-7
ROWS = [
    ("Learning Targets", "learning_targets", 'Start with "I can…"'),
    ("Standards", "standards", "Standard:"),
    ("ACT Alignment", "act_alignment", "ACT Equivalent:"),
    ("Engagement Strategy", "engagement_strategy", ""),
    ("Lesson", "lesson", "Do Now / During / Assessment"),
]


def shade(cell, hex_fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcpr.append(shd)


def set_width(cell, dxa):
    cell.width = Twips(dxa)
    tcpr = cell._tc.get_or_add_tcPr()
    tcw = OxmlElement("w:tcW")
    tcw.set(qn("w:w"), str(dxa))
    tcw.set(qn("w:type"), "dxa")
    tcpr.append(tcw)


def write_cell(cell, blocks, font="Calibri", size=11, bold=False,
               color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT):
    """blocks: list of (text, bold, size) or a plain string (split on newlines)."""
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    if isinstance(blocks, str):
        lines = blocks.split("\n")
        blocks = [(ln, bold, size) for ln in lines]
    first = True
    for text, b, sz in blocks:
        if not first:
            para = cell.add_paragraph()
            para.alignment = align
        run = para.add_run(text)
        run.font.name = font
        run.font.size = Pt(sz)
        run.font.bold = b
        run.font.color.rgb = None  # default black
        first = False


def fixed_layout(table):
    tbl = table._tbl
    tblpr = tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblpr.append(layout)


def build(data, out_path):
    doc = Document()

    # Landscape letter, 0.5" margins
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Inches(11)
    sec.page_height = Inches(8.5)
    sec.top_margin = sec.bottom_margin = Inches(0.5)
    sec.left_margin = sec.right_margin = Inches(0.5)

    days = {d["name"]: d for d in data.get("days", [])}

    table = doc.add_table(rows=0, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    fixed_layout(table)

    # Row 1: merged header — teacher | grade/subject | week of
    r1 = table.add_row().cells
    r1[0].merge(r1[5])
    header = (f"Teacher: {data.get('teacher','')}    |    "
              f"{data.get('grade_subject','')}    |    "
              f"Week of: {data.get('week_of','')}")
    write_cell(r1[0], [(header, True, 11)], font="Arial", bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    shade(r1[0], BLUE_HDR)

    # Row 2: label + day names
    r2 = table.add_row().cells
    set_width(r2[0], COL_DXA[0])
    write_cell(r2[0], [("Lesson Plan Components", True, 11)], font="Arial", bold=True)
    shade(r2[0], BLUE_HDR)
    for i, name in enumerate(DAY_NAMES, start=1):
        set_width(r2[i], COL_DXA[i])
        write_cell(r2[i], [(name, True, 11)], font="Arial", bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        shade(r2[i], BLUE_HDR)

    # Rows 3-8: label col + content
    for label, key, guidance in ROWS:
        row = table.add_row().cells
        set_width(row[0], COL_DXA[0])
        label_blocks = [(label, True, 11)]
        if guidance:
            label_blocks.append((guidance, False, 9))
        write_cell(row[0], label_blocks, font="Calibri")
        shade(row[0], BLUE_LBL)
        for i, name in enumerate(DAY_NAMES, start=1):
            set_width(row[i], COL_DXA[i])
            d = days.get(name, {})
            if d.get("no_school"):
                # Only stamp "No School" once, in the first content row
                if key == "curriculum":
                    write_cell(row[i], [("No School", True, 11)], font="Arial",
                               bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
                else:
                    write_cell(row[i], "")
                shade(row[i], WHITE)
                continue
            text = str(d.get(key, "")).strip()
            
            # Format Engagement Strategy as Bold
            is_engagement = (key == "engagement_strategy" and bool(text))
            
            write_cell(row[i], [(text, is_engagement, 11)], font="Calibri")
            shade(row[i], WHITE)

    doc.save(out_path)
    print(f"Wrote {out_path}")


def main():
    if len(sys.argv) < 3:
        print("Usage: python3 build_lesson_plan.py <week.json> <output.docx>")
        sys.exit(1)
    with open(sys.argv[1], encoding="utf-8") as fh:
        data = json.load(fh)
    build(data, sys.argv[2])


if __name__ == "__main__":
    main()
