"""Ad-hoc checks for template_intake.py's deterministic stages — extraction,
structural sanity checks, and LLM-output cross-validation — run directly
with `python test_template_intake.py`, matching this repo's existing
test_db.py/test_courses.py style rather than a pytest suite.

Deliberately does NOT touch Postgres or OpenAI: everything here is pure
function behavior on synthetic docx files and hand-built dicts, so it can
run anywhere, including CI with no .env at all.
"""
import sys
import tempfile
from pathlib import Path

import docx

sys.path.insert(0, ".")
from backend import template_intake as ti

failures = []


def check(label, condition):
    status = "ok" if condition else "FAIL"
    print(f"[{status}] {label}")
    if not condition:
        failures.append(label)


def make_docx(build_fn) -> Path:
    d = docx.Document()
    build_fn(d)
    fd = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    fd.close()
    d.save(fd.name)
    return Path(fd.name)


# ---------------------------------------------------------------------------
# 1. A well-formed template: headings, a clean table, some bold text.
# ---------------------------------------------------------------------------


def build_good(d):
    d.add_heading("Weekly Lesson Plan Template", level=0)
    d.add_heading("Header Information", level=1)
    t = d.add_table(rows=2, cols=3)
    t.rows[0].cells[0].text = "Day"
    t.rows[0].cells[1].text = "Standard"
    t.rows[0].cells[2].text = "Objective"
    d.add_heading("Instructional Notes", level=1)
    p = d.add_paragraph("Fill in ")
    p.add_run("daily objectives").bold = True
    p.add_run(" here.")


good_path = make_docx(build_good)
good_structure = ti._extract_docx_structure(good_path)
good_findings = ti._run_structural_checks(good_structure)
check(
    "well-formed template: no error-severity findings",
    not any(f.severity == "error" for f in good_findings),
)
check("well-formed template: 3 headings extracted", len(good_structure["headings"]) == 3)
check("well-formed template: 1 clean table extracted", len(good_structure["tables"]) == 1)
good_path.unlink()


# ---------------------------------------------------------------------------
# 2. No visible structure at all: plain prose, no headings/tables/emphasis.
# ---------------------------------------------------------------------------


def build_structureless(d):
    d.add_paragraph("This is just some plain text with nothing distinguishing about it at all.")
    d.add_paragraph("Another plain paragraph, still nothing a builder script could hook into.")


bad_path = make_docx(build_structureless)
bad_structure = ti._extract_docx_structure(bad_path)
bad_findings = ti._run_structural_checks(bad_structure)
check(
    "structureless doc: no_visible_structure error fires",
    any(f.check_name == "no_visible_structure" and f.severity == "error" for f in bad_findings),
)
bad_path.unlink()


# ---------------------------------------------------------------------------
# 3. Ragged table: rows with inconsistent cell counts should warn.
# ---------------------------------------------------------------------------


def build_ragged(d):
    d.add_heading("Weekly Lesson Plan Template", level=1)
    d.add_paragraph("Fill in each day's objective, standard, and materials below.")
    t = d.add_table(rows=2, cols=3)
    t.rows[0].cells[0].text = "Day"
    t.rows[1].cells[0].text = "Monday"
    t.rows[1].cells[1].text = "Objective text goes here"
    t.rows[1].cells[2].text = "Standard code"
    row = t.rows[0]._tr
    row.remove(row.tc_lst[-1])  # manually drop a cell to make row 0 narrower than row 1


ragged_path = make_docx(build_ragged)
ragged_structure = ti._extract_docx_structure(ragged_path)
ragged_findings = ti._run_structural_checks(ragged_structure)
check(
    "ragged table: ragged_table warning fires",
    any(f.check_name == "ragged_table" for f in ragged_findings),
)
ragged_path.unlink()


# ---------------------------------------------------------------------------
# 4. Empty document: near-zero text should be caught before anything else.
# ---------------------------------------------------------------------------


def build_empty(d):
    d.add_paragraph("")


empty_path = make_docx(build_empty)
empty_structure = ti._extract_docx_structure(empty_path)
empty_findings = ti._run_structural_checks(empty_structure)
check(
    "empty doc: empty_document error fires",
    any(f.check_name == "empty_document" and f.severity == "error" for f in empty_findings),
)
empty_path.unlink()


# ---------------------------------------------------------------------------
# 5. Upload-boundary validation: extension must match magic bytes.
# ---------------------------------------------------------------------------

fake_docx = tempfile.NamedTemporaryFile(suffix=".tmp", delete=False)
fake_docx.write(b"this is not a real docx file, just plain bytes")
fake_docx.close()
fake_docx_path = Path(fake_docx.name)

try:
    ti.validate_upload("template.docx", fake_docx_path)
    check("magic-byte mismatch is rejected", False)
except Exception as e:
    check(f"magic-byte mismatch is rejected ({type(e).__name__})", type(e).__name__ == "AppError")
fake_docx_path.unlink()

real_docx_path = make_docx(build_good)
ext = ti.validate_upload("template.docx", real_docx_path)
check("real docx passes validate_upload", ext == ".docx")
real_docx_path.unlink()


# ---------------------------------------------------------------------------
# 6. Anti-hallucination cross-validation: a grounded claim survives, a
#    fabricated one is dropped and produces a finding.
# ---------------------------------------------------------------------------

summary_text = "Headings:\n  [Heading 1] Learning Objectives\nTables:\n  Table 0: header row: ['Day', 'Standard']"

fake_analysis = {
    "sections": [
        {
            "name": "Objectives",
            "description": "Daily learning objectives",
            "source_evidence": "Learning Objectives",
            "repeats_per_entry": False,
        },
        {
            "name": "Fabricated section",
            "description": "Something the model made up",
            "source_evidence": "Standards Alignment Rubric",  # not present in summary_text at all
            "repeats_per_entry": False,
        },
    ],
    "unclear_or_ambiguous": [],
    "overall_confidence": 0.8,
    "recommended_for_auto_use": True,
}

verified, verify_findings = ti._cross_validate_llm_output(fake_analysis, summary_text)
check("grounded section survives cross-validation", len(verified["sections"]) == 1)
check("grounded section is the correct one", verified["sections"][0]["name"] == "Objectives")
check(
    "fabricated section produces an unverified_claim finding",
    any(f.check_name == "unverified_claim" for f in verify_findings),
)
check(
    "recommend-auto-use + dropped section flags a confidence_mismatch",
    any(f.check_name == "confidence_mismatch" for f in verify_findings),
)


# ---------------------------------------------------------------------------
# 7. Coverage check: a heading nobody's section (or unclear list) mentions
#    should be flagged — the complement to hallucination-catching.
# ---------------------------------------------------------------------------

covered_structure = {
    "format": "docx",
    "headings": [
        {"text": "Learning Objectives", "style": "Heading 1", "level": 1},
        {"text": "Materials Needed", "style": "Heading 1", "level": 1},
    ],
    "tables": [],
}
analysis_missing_one = {
    "sections": [
        {
            "name": "Objectives",
            "description": "Daily objectives",
            "source_evidence": "Learning Objectives",
            "repeats_per_entry": False,
        }
    ],
    "unclear_or_ambiguous": [],
}
coverage_findings = ti._check_coverage(covered_structure, analysis_missing_one)
check(
    "coverage check flags the unmentioned heading",
    any(f.check_name == "uncovered_heading" and "Materials Needed" in f.message for f in coverage_findings),
)

analysis_fully_covered = {
    "sections": [
        {
            "name": "Objectives",
            "description": "Daily objectives",
            "source_evidence": "Learning Objectives",
            "repeats_per_entry": False,
        }
    ],
    "unclear_or_ambiguous": ["Materials Needed"],
}
check(
    "coverage check clears when the rest is listed as unclear",
    not ti._check_coverage(covered_structure, analysis_fully_covered),
)


# ---------------------------------------------------------------------------
# 8. Second-pass verifier: rejects a section the model itself flags as
#    inaccurate, and fails open (keeps sections) if the verifier errors.
# ---------------------------------------------------------------------------

two_section_analysis = {
    "sections": [
        {
            "name": "Objectives",
            "description": "Daily objectives",
            "source_evidence": "Learning Objectives",
            "repeats_per_entry": False,
        },
        {
            "name": "Materials",
            "description": "Materials needed for the lesson",
            "source_evidence": "Materials Needed",
            "repeats_per_entry": False,
        },
    ],
    "unclear_or_ambiguous": [],
}

original_verify = ti.llm.verify_template_sections
ti.llm.verify_template_sections = lambda user_id, summary, sections: {
    "verdicts": [
        {"name": "Objectives", "accurate": True, "reason": "matches"},
        {"name": "Materials", "accurate": False, "reason": "description overstates the evidence"},
    ]
}
verified2, second_pass_findings = ti._apply_second_pass_verifier("test_user", two_section_analysis, "irrelevant")
check("second-pass verifier keeps the accurate section", len(verified2["sections"]) == 1)
check("second-pass verifier drops the rejected section", verified2["sections"][0]["name"] == "Objectives")
check(
    "second-pass rejection produces a finding",
    any(f.check_name == "second_pass_rejected" for f in second_pass_findings),
)

ti.llm.verify_template_sections = lambda *a, **kw: (_ for _ in ()).throw(RuntimeError("network down"))
verified3, fail_open_findings = ti._apply_second_pass_verifier("test_user", two_section_analysis, "irrelevant")
check("second-pass verifier fails open (keeps both sections)", len(verified3["sections"]) == 2)
check(
    "second-pass failure produces an info finding, not an error",
    any(f.check_name == "second_pass_unavailable" and f.severity == "info" for f in fail_open_findings),
)
ti.llm.verify_template_sections = original_verify


# ---------------------------------------------------------------------------
# 9. OCR fallback: a scanned (image-only, no text layer) PDF should still
#    yield readable structure when tesseract is installed locally.
# ---------------------------------------------------------------------------

if not ti._tesseract_available():
    print("[skip] OCR fallback: tesseract not installed in this environment")
else:
    from PIL import Image, ImageDraw

    img = Image.new("RGB", (1700, 2200), "white")
    drawer = ImageDraw.Draw(img)
    drawer.text((100, 100), "Weekly Lesson Plan Template", fill="black")
    drawer.text((100, 200), "Fill in the objective, standard, and materials for each day of the week.", fill="black")
    scanned_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    scanned_pdf.close()
    img.save(scanned_pdf.name, "PDF", resolution=150.0)
    scanned_path = Path(scanned_pdf.name)

    scanned_structure = ti._extract_pdf_structure(scanned_path)
    check("OCR fallback triggers on an image-only PDF", scanned_structure["ocr_status"] == "succeeded")
    check(
        "OCR recovers the rendered text",
        "Weekly Lesson Plan Template".lower() in (scanned_structure["ocr_text"] or "").lower(),
    )
    scanned_findings = ti._run_structural_checks(scanned_structure)
    check(
        "OCR'd doc gets an ocr_used warning, not a hard error",
        any(f.check_name == "ocr_used" for f in scanned_findings)
        and not any(f.severity == "error" for f in scanned_findings),
    )
    scanned_path.unlink()


# ---------------------------------------------------------------------------
# 10. Auto-activation bar: only a genuinely unanimous result should pass.
# ---------------------------------------------------------------------------

clean_result = {
    "status": "analyzed",
    "findings": [],
    "analysis": {
        "sections": [{"name": "Objectives", "description": "d", "source_evidence": "e", "repeats_per_entry": False}],
        "overall_confidence": 0.95,
        "recommended_for_auto_use": True,
    },
}
check("auto-activation bar: clean result passes", ti._meets_auto_activation_bar(clean_result))

warned_result = dict(clean_result, status="analyzed_with_warnings")
check("auto-activation bar: any non-'analyzed' status fails", not ti._meets_auto_activation_bar(warned_result))

with_info_finding = dict(clean_result, findings=[{"stage": "x", "check_name": "y", "severity": "info", "message": "m"}])
check(
    "auto-activation bar: even an info-severity finding fails it (not just errors/warnings)",
    not ti._meets_auto_activation_bar(with_info_finding),
)

low_confidence = dict(clean_result, analysis=dict(clean_result["analysis"], overall_confidence=0.5))
check("auto-activation bar: low confidence fails", not ti._meets_auto_activation_bar(low_confidence))

not_recommended = dict(clean_result, analysis=dict(clean_result["analysis"], recommended_for_auto_use=False))
check("auto-activation bar: model not recommending auto-use fails", not ti._meets_auto_activation_bar(not_recommended))

no_sections = dict(clean_result, analysis=dict(clean_result["analysis"], sections=[]))
check("auto-activation bar: empty section list fails", not ti._meets_auto_activation_bar(no_sections))


print()
if failures:
    print(f"{len(failures)} check(s) failed: {failures}")
    sys.exit(1)
print("All checks passed.")
